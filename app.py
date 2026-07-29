import os
import json
import re
import io
import fitz  # PyMuPDF
import openpyxl
import docx
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask import Flask, request, jsonify, render_template, send_file

import pytesseract
from PIL import Image

import sys
import webbrowser
import threading

# ─── Database Adapter ─────────────────────────────────────────────────────────
DB_ADAPTER = None

def init_db_adapter():
    global DB_ADAPTER
    backend = os.environ.get("DB_BACKEND", "local").lower()
    if backend == "supabase":
        from db_adapter import get_adapter
        DB_ADAPTER = get_adapter(
            backend="supabase",
            supabase_url=os.environ.get("SUPABASE_URL", ""),
            supabase_key=os.environ.get("SUPABASE_KEY", "")
        )
        print(f"Database backend: Supabase ({os.environ.get('SUPABASE_URL', '')[:30]}...)")
    else:
        from db_adapter import get_adapter
        db_path = os.environ.get("DB_PATH", os.path.join(APP_DIR, "DATABASE.xlsx"))
        DB_ADAPTER = get_adapter(backend="local", db_path=db_path)
        print(f"Database backend: Excel locale ({DB_PATH})")

# ─── Portable & PyInstaller Base Paths ─────────────────────────────────────────
if getattr(sys, 'frozen', False):
    BASE_DIR = sys._MEIPASS
    APP_DIR  = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    APP_DIR  = BASE_DIR

# ─── Tesseract Path Resolution ────────────────────────────────────────────────
LOCAL_TESS  = os.path.join(APP_DIR, "Tesseract-OCR", "tesseract.exe")
BUNDLE_TESS = os.path.join(BASE_DIR, "Tesseract-OCR", "tesseract.exe")
USER_TESS   = r"C:\Users\d.delisa.CONTEA\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"

if os.path.exists(LOCAL_TESS):
    pytesseract.pytesseract.tesseract_cmd = LOCAL_TESS
    os.environ['TESSDATA_PREFIX'] = os.path.join(APP_DIR, "Tesseract-OCR", "tessdata")
elif os.path.exists(BUNDLE_TESS):
    pytesseract.pytesseract.tesseract_cmd = BUNDLE_TESS
    os.environ['TESSDATA_PREFIX'] = os.path.join(BASE_DIR, "Tesseract-OCR", "tessdata")
elif os.path.exists(USER_TESS):
    pytesseract.pytesseract.tesseract_cmd = USER_TESS

# ─── Flask App with PyInstaller Resource Paths ────────────────────────────────
template_folder = os.path.join(BASE_DIR, 'templates')
static_folder   = os.path.join(BASE_DIR, 'static')
app = Flask(__name__, template_folder=template_folder, static_folder=static_folder)

# ─── CORS ─────────────────────────────────────────────────────────────────────
from flask_cors import CORS
CORS(app, resources={r"/api/*": {"origins": "*"}})

# ─── Database Path Resolution (solo per compatibilità locale) ─────────────────
DOCS_DIR      = os.path.join(os.path.expanduser("~"), "Documents", "CONTEA DVR Analyzer")
DOCS_DB       = os.path.join(DOCS_DIR, "DATABASE.xlsx")
BUNDLED_DB    = os.path.join(APP_DIR,  "DATABASE.xlsx")
BUNDLED_DB_2  = os.path.join(BASE_DIR, "DATABASE.xlsx")

def get_db_path():
    """Restituisce sempre il percorso del DATABASE.xlsx in Documenti (scrivibile)."""
    os.makedirs(DOCS_DIR, exist_ok=True)
    if not os.path.exists(DOCS_DB):
        for src in [BUNDLED_DB, BUNDLED_DB_2]:
            if os.path.exists(src):
                import shutil
                shutil.copy2(src, DOCS_DB)
                print(f"DATABASE.xlsx copiato in: {DOCS_DB}")
                break
    return DOCS_DB

DB_PATH = get_db_path()
init_db_adapter()

# ─── Database helpers (compatibilità) ─────────────────────────────────────────

def normalize_key(k):
    """Normalise a checklist key for fuzzy matching."""
    if not k:
        return ""
    k = str(k).lower().strip()
    k = re.sub(r'\(.*?\)', '', k)
    k = re.sub(r'[^a-z0-9\s\n]', ' ', k)
    return " ".join(k.split())

def read_checklist_items():
    """Legge le voci della checklist dal database (via adapter)."""
    if DB_ADAPTER is None:
        return []
    items = DB_ADAPTER.get_checklist_items()
    return [
        {
            'category': it.get('category', ""),
            'item': it.get('item', ""),
            'key': it.get('key', ""),
            'is_checkbox': it.get('is_checkbox', True)
        }
        for it in items
    ]

def get_row_images_map():
    """Per Supabase non serve più: le immagini sono URL già pronte."""
    if DB_ADAPTER is None:
        return {}
    # ExcelAdapter mantiene la cache interna
    if hasattr(DB_ADAPTER, '_get_images_map'):
        return DB_ADAPTER._get_images_map()
    return {}

def match_risks(checked_keys):
    """Associa le chiavi rilevate alle righe DVR (via adapter)."""
    if DB_ADAPTER is None:
        return []
    return DB_ADAPTER.match_risks(checked_keys)

def get_all_dvr_risks():
    """Legge tutti i rischi DVR ordinati per colonna H (via adapter)."""
    if DB_ADAPTER is None:
        return []
    return DB_ADAPTER.get_all_dvr_risks()

def get_structured_checklist():
    """Legge le voci della checklist e le raggruppa per Categoria."""
    cl_items = read_checklist_items()
    categories = {}
    for item in cl_items:
        cat = item['category'] or "GENERALE"
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(item)
    return [{'category': cat_name, 'items': items} for cat_name, items in categories.items()]

# ─── Analisi PDF locale (senza API) ──────────────────────────────────────────

# Marcatori visivi di checkbox spuntato (testo/unicode)
CHECKBOX_MARKERS = ['[x]', '[X]', '☑', '☒', '✔', '✓', '✗', 'þ', 'R']

def analyze_pdf_local(pdf_bytes, checklist_items):
    """
    Analizza il PDF in locale senza alcuna API esterna.

    Strategia in 3 livelli:
      1. Legge i widget di tipo checkbox dal PDF (funziona per PDF digitali con form).
      2. Estrae il testo grezzo e cerca marcatori di spunta vicino ai nomi degli item
         (funziona per PDF di testo, checklist compilate digitalmente ma senza form fields).
      3. Se non rileva testo, esegue l'OCR (Tesseract) sulle pagine convertite in immagini
         e associa spazialmente i checkmark rilevati nelle colonne di spunta.

    Restituisce la lista delle chiavi rilevate come spuntate.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    checkbox_items = [item for item in checklist_items if item['is_checkbox']]
    checked_keys = []

def analyze_pdf_local(pdf_bytes, checklist_items, session_id=None):
    """
    Analizza il PDF in locale senza alcuna API esterna e riporta il progresso in tempo reale.
    """
    def update_prog(pct, msg):
        if session_id:
            ANALYSIS_PROGRESS[session_id] = {"pct": pct, "msg": msg}
            print(f"[Progress {pct}%] {msg}")

    update_prog(10, "Conversione del documento PDF...")
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    num_pages = len(doc)
    checkbox_items = [item for item in checklist_items if item['is_checkbox']]
    checked_keys = []

    # ── Livello 1: form fields ────────────────────────────────────────────────
    update_prog(20, "Verifica moduli ed elementi interattivi (Livello 1)...")
    widget_names_checked = set()
    for page in doc:
        for widget in page.widgets():
            if widget.field_type == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                val = widget.field_value
                if val in (True, 'Yes', 'On', 'true', 'yes', 'on', '1'):
                    fname = (widget.field_name or "").strip()
                    if fname:
                        widget_names_checked.add(fname)

    if widget_names_checked:
        for item in checkbox_items:
            key      = item['key']
            norm_key = normalize_key(key)
            for wname in widget_names_checked:
                norm_w = normalize_key(wname)
                if key == wname or norm_key == norm_w or norm_key in norm_w or norm_w in norm_key:
                    if key not in checked_keys:
                        checked_keys.append(key)
                    break
        if checked_keys:
            update_prog(100, f"Rilevati {len(checked_keys)} flag da moduli PDF!")
            return checked_keys

    # ── Livello 2: analisi testo ──────────────────────────────────────────────
    update_prog(30, "Analisi testo nativo e marcatori visivi (Livello 2)...")
    full_text = ""
    for page in doc:
        full_text += page.get_text() + "\n"

    if len(full_text.strip()) > 100:
        for item in checkbox_items:
            key     = item['key']
            pattern = re.escape(key)
            for line in full_text.splitlines():
                line_stripped = line.strip()
                if not line_stripped:
                    continue
                has_marker = any(m in line for m in CHECKBOX_MARKERS)
                key_found  = bool(re.search(pattern, line, re.IGNORECASE))
                if has_marker and key_found:
                    if key not in checked_keys:
                        checked_keys.append(key)
                    break
        if checked_keys:
            update_prog(100, f"Rilevati {len(checked_keys)} flag da testo nativo!")
            return checked_keys

    # ── Livello 3: OCR spaziale (Tesseract) con preprocessing ──────────────────
    update_prog(40, "Avvio motore OCR Tesseract per scansione cartacea (Livello 3)...")
    
    try:
        pytesseract.get_tesseract_version()
    except Exception as e:
        print("Tesseract OCR non disponibile:", e)
        return []

    # Preprocessing OpenCV opzionale per migliorare velocità e accuratezza OCR
    def preprocess_for_ocr(pil_img):
        nx, ny = pil_img.size
        # Ridimensiona se troppo grande (max 2000px larghezza)
        if nx > 2000:
            ratio = 2000 / nx
            nx = 2000
            ny = int(ny * ratio)
            pil_img = pil_img.resize((nx, ny), Image.LANCZOS)

        # Converti in array numpy per OpenCV se disponibile
        try:
            import cv2
            import numpy as np
            arr = np.array(pil_img.convert("RGB"))
            gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
            # Denoising leggero
            denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
            # Binarizzazione adattiva per checkmark e testo
            binary = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY, 11, 2
            )
            pil_img = Image.fromarray(binary)
        except ImportError:
            pass
        return pil_img

    lang = 'ita'
    try:
        available_langs = pytesseract.get_languages()
        if 'ita' not in available_langs:
            lang = 'eng'
    except Exception:
        lang = 'ita'

    for page_idx, page in enumerate(doc):
        start_pct = 40 + int((page_idx / num_pages) * 45)
        update_prog(start_pct, f"Esecuzione OCR Pagina {page_idx + 1} di {num_pages}...")
        
        pix = page.get_pixmap(dpi=150)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        img = preprocess_for_ocr(img)
        
        try:
            data_str = pytesseract.image_to_data(img, lang=lang)
            lines = data_str.splitlines()
            if not lines:
                continue
            
            header = lines[0].split("\t")
            words = []
            for line in lines[1:]:
                parts = line.split("\t")
                if len(parts) == len(header):
                    text = parts[-1].strip()
                    if text:
                        words.append({
                            "left": int(parts[header.index("left")]),
                            "top": int(parts[header.index("top")]),
                            "width": int(parts[header.index("width")]),
                            "height": int(parts[header.index("height")]),
                            "text": text
                        })
            
            page_item_coords = []
            for item in checkbox_items:
                item_words = re.sub(r"[^a-zA-Z0-9\s]", " ", item["item"]).lower().split()
                if not item_words:
                    continue
                    
                matching_words = []
                for iw in item_words:
                    if len(iw) <= 2:
                        continue
                    for w in words:
                        if iw in w["text"].lower():
                            matching_words.append(w)
                            
                if matching_words:
                    y_votes = {}
                    for w in matching_words:
                        top = w["top"]
                        found = False
                        for yt in y_votes:
                            if abs(yt - top) < 12:
                                y_votes[yt].append(w)
                                found = True
                                break
                        if not found:
                            y_votes[top] = [w]
                    
                    best_y = None
                    max_votes = 0
                    for yt, vlist in y_votes.items():
                        if len(vlist) > max_votes:
                            max_votes = len(vlist)
                            best_y = yt
                    
                    avg_left = sum(w["left"] for w in y_votes[best_y]) / len(y_votes[best_y])
                    column = 1 if avg_left < 500 else 2
                    
                    page_item_coords.append({
                        "key": item["key"],
                        "y": best_y,
                        "column": column
                    })
            
            detected_marks = []
            for w in words:
                if len(w["text"]) > 4:
                    continue
                
                is_mark = False
                column = None
                if 300 <= w["left"] <= 420:
                    is_mark = True
                    column = 1
                elif 880 <= w["left"] <= 975:
                    is_mark = True
                    column = 2
                    
                if is_mark:
                    detected_marks.append({
                        "y": w["top"],
                        "column": column
                    })
            
            for mark in detected_marks:
                best_match = None
                min_dist = 20
                for pic in page_item_coords:
                    if pic["column"] == mark["column"]:
                        dist = abs(pic["y"] - mark["y"])
                        if dist < min_dist:
                            min_dist = dist
                            best_match = pic
                
                if best_match:
                    key = best_match["key"]
                    if key not in checked_keys:
                        checked_keys.append(key)
                        
        except Exception as ocr_err:
            print(f"Errore OCR pagina {page_idx+1}:", ocr_err)

    update_prog(90, "Mappatura risposte e riordinamento criticità...")
    return checked_keys

# ─── Word helpers ─────────────────────────────────────────────────────────────

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr  = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'),    str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr      = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '8')  # 1pt solid line
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Solid black
        tblBorders.append(border)
    tblPr.append(tblBorders)

# ─── Flask routes ─────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/progress/<session_id>', methods=['GET'])
def get_progress(session_id):
    prog = ANALYSIS_PROGRESS.get(session_id, {"pct": 0, "msg": "Inizializzazione..."})
    return jsonify(prog)

@app.route('/api/checklist_structure', methods=['GET'])
def checklist_structure():
    try:
        data = get_structured_checklist()
        return jsonify({"success": True, "categories": data})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/database_risks', methods=['GET'])
def database_risks():
    try:
        risks = get_all_dvr_risks()
        return jsonify({"success": True, "count": len(risks), "risks": risks})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/add_database_risk', methods=['POST'])
def add_database_risk():
    global _excel_row_images_cache
    try:
        data = request.json or {}
        luogo = (data.get("luogo") or "").strip()
        rischio = (data.get("rischio") or "").strip()
        azione = (data.get("azione") or "").strip()
        entro = (data.get("entro") or "").strip()
        raw_ordine = data.get("ordine")
        
        try:
            ordine = int(raw_ordine) if raw_ordine is not None else 999999
        except (ValueError, TypeError):
            ordine = 999999

        if not rischio and not luogo:
            return jsonify({"error": "Inserire almeno il Luogo o il Fattore di Rischio."}), 400

        if DB_ADAPTER is None:
            return jsonify({"error": "Database non inizializzato."}), 500

        risk = {
            "luogo": luogo,
            "rischio": rischio,
            "azione": azione,
            "entro": entro,
            "ordine": ordine,
            "image": data.get("image")
        }
        result = DB_ADAPTER.add_database_risk(risk)
        if isinstance(DB_ADAPTER, ExcelAdapter):
            _excel_row_images_cache = None
        return jsonify(result)

    except Exception as e:
        print("Errore salvataggio riga:", e)
        return jsonify({"error": str(e)}), 500

@app.route('/api/match_keys', methods=['POST'])
def match_selected_keys():
    try:
        data = request.json or {}
        keys = data.get("keys", [])
        risks = match_risks(keys)
        return jsonify({"success": True, "matched_count": len(risks), "risks": risks})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/analyze', methods=['POST'])
def analyze_pdf():
    if 'file' not in request.files:
        return jsonify({"error": "Nessun file caricato"}), 400

    file = request.files['file']
    session_id = request.form.get('session_id') or 'default'
    if not file.filename:
        return jsonify({"error": "Nome file vuoto"}), 400

    try:
        pdf_bytes = file.read()

        cl_items = read_checklist_items()
        if not cl_items:
            return jsonify({"error": f"Impossibile leggere il database: {DB_PATH}"}), 500

        detected_keys = analyze_pdf_local(pdf_bytes, cl_items, session_id=session_id)
        risks = match_risks(detected_keys)
        ANALYSIS_PROGRESS[session_id] = {"pct": 100, "msg": "Completato!"}

        return jsonify({
            "success":        True,
            "detected_keys":  detected_keys,
            "detected_count": len(detected_keys),
            "matched_count":  len(risks),
            "risks":          risks
        })

    except Exception as e:
        print("Errore analisi:", e)
        ANALYSIS_PROGRESS[session_id] = {"pct": 0, "msg": f"Errore: {e}"}
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate_docx', methods=['POST'])
def generate_docx():
    import base64
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data  = request.json or {}
    risks = data.get("risks", [])

    try:
        def get_order_val(r):
            val = r.get("ordine", 999999)
            try:
                return int(val) if val is not None else 999999
            except (ValueError, TypeError):
                return 999999

        risks.sort(key=get_order_val)

        document = docx.Document()

        section          = document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation  = WD_ORIENT.LANDSCAPE
        section.page_width   = new_width
        section.page_height  = new_height

        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

        title_p = document.add_paragraph()
        run = title_p.add_run("CRITICITÀ RILEVATE DAL SOPRALLUOGO")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Arial"
        title_p.paragraph_format.space_after = Pt(12)

        table = document.add_table(rows=1, cols=5)
        table.autofit = False
        set_table_borders(table)

        headers = [
            "Luogo di lavoro/Impianto",
            "Fattore di rischio / Anomalia riscontrata",
            "Azione migliorativa proposta",
            "Spazio Foto",
            "Da attuarsi entro"
        ]
        widths = [Inches(1.8), Inches(1.8), Inches(4.0), Inches(1.8), Inches(1.0)]

        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = text
            hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_background(hdr_cells[idx], "E6E6E6")
            set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=120, right=120)
            p = hdr_cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.name = "Arial"

        for idx, cell in enumerate(hdr_cells):
            cell.width = widths[idx]

        for risk in risks:
            row_cells = table.add_row().cells
            row_cells[0].text = risk.get("luogo",   "")
            row_cells[1].text = risk.get("rischio", "")
            row_cells[2].text = risk.get("azione",  "")
            row_cells[3].text = ""
            row_cells[4].text = risk.get("entro",   "")

            # Allinea rigorosamente ogni cella IN ALTO (VALIGN = TOP) come richiesto
            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

            # Inserimento Immagine Proporzionale (Intera, ridimensionata senza alcuna distorsione)
            img_data = risk.get("image")
            if img_data:
                try:
                    if img_data.startswith("data:"):
                        header_str, b64_str = img_data.split(",", 1)
                        img_bytes = base64.b64decode(b64_str)
                    else:
                        img_bytes = base64.b64decode(img_data)

                    pil_img = Image.open(io.BytesIO(img_bytes))
                    orig_w, orig_h = pil_img.size

                    if orig_w > 0 and orig_h > 0:
                        aspect = orig_h / orig_w
                        max_w = 1.35  # Max larghezza in pollici per non spaginare
                        max_h = 1.05  # Max altezza in pollici per rimanere compatta ed intera

                        if (max_w * aspect) > max_h:
                            target_h = max_h
                            target_w = target_h / aspect
                        else:
                            target_w = max_w
                            target_h = target_w * aspect

                        img_stream = io.BytesIO(img_bytes)
                        cell_p = row_cells[3].paragraphs[0]
                        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell_p.paragraph_format.space_after = Pt(0)
                        run_img = cell_p.add_run()
                        run_img.add_picture(img_stream, width=Inches(target_w), height=Inches(target_h))
                except Exception as img_err:
                    print("Errore inserimento immagine nel Word:", img_err)

            for idx, cell in enumerate(row_cells):
                cell.width = widths[idx]
                set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Arial"

        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name="criticita_dvr_rilevate.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print("Errore generazione DOCX:", e)
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    def open_browser():
        try:
            webbrowser.open("http://127.0.0.1:5000")
        except Exception:
            pass

    threading.Timer(1.2, open_browser).start()
    print("Avvio del server DVR Checklist Analyzer su http://127.0.0.1:5000...")
    app.run(host='127.0.0.1', port=5000, debug=False)
